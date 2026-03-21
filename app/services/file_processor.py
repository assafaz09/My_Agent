import os
import asyncio
import aiofiles
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
import speech_recognition as sr
import logging
from datetime import datetime

from app.core.config import settings
from app.services.openai_service import OpenAIService
from app.services.qdrant_service import QdrantService
from app.models.knowledge import KnowledgeDocument, KnowledgeChunk

logger = logging.getLogger(__name__)

class FileProcessor:
    def __init__(self):
        self.openai_service = OpenAIService()
        self.qdrant_service = QdrantService()
        
        # Initialize speech recognition
        self.recognizer = sr.Recognizer()
    
    async def process_file_async(
        self, 
        upload_id: str, 
        filename: str, 
        file_path: str, 
        file_type: str
    ):
        """Process file in background"""
        try:
            # Update status to processing
            await self._update_upload_status(upload_id, "processing", 0.1)
            
            # Extract text from file
            text = await self._extract_text(file_path, file_type)
            await self._update_upload_status(upload_id, "processing", 0.4)
            
            # Extract knowledge using OpenAI
            knowledge_data = await self.openai_service.extract_knowledge(text, file_type)
            await self._update_upload_status(upload_id, "processing", 0.6)
            
            # Create document
            document = KnowledgeDocument(
                document_id=upload_id,
                filename=filename,
                file_type=file_type,
                content=text,
                metadata={
                    "extracted_knowledge": knowledge_data,
                    "file_size": os.path.getsize(file_path),
                    "processed_at": datetime.now().isoformat()
                },
                created_at=datetime.now()
            )
            
            # Split into chunks and generate embeddings
            chunks = await self._create_chunks_with_embeddings(document)
            await self._update_upload_status(upload_id, "processing", 0.8)
            
            # Store in Qdrant
            try:
                await self.qdrant_service.add_chunks_with_embeddings(chunks)
                await self._update_upload_status(upload_id, "completed", 1.0)
            except Exception as e:
                logger.error(f"Error storing in Qdrant: {e}")
                await self._update_upload_status(upload_id, "error", 0.0, f"Qdrant storage error: {str(e)}")
                return
            
            logger.info(f"Successfully processed file: {filename}")
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            await self._update_upload_status(upload_id, "error", 0.0, str(e))
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file types"""
        try:
            if file_type == "txt":
                return await self._extract_from_txt(file_path)
            elif file_type == "pdf":
                return await self._extract_from_pdf(file_path)
            elif file_type == "docx":
                return await self._extract_from_docx(file_path)
            elif file_type in ["jpg", "jpeg", "png"]:
                return await self._extract_from_image(file_path)
            elif file_type in ["wav", "mp3", "m4a"]:
                return await self._extract_from_audio(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    async def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logger.error(f"OCR error for {file_path}: {e}")
            return ""
    
    async def _extract_from_audio(self, file_path: str) -> str:
        """Extract text from audio file using speech recognition"""
        try:
            with sr.AudioFile(file_path) as source:
                audio = self.recognizer.record(source)
                
            # Try Google Speech Recognition first
            try:
                text = self.recognizer.recognize_google(audio)
                return text
            except sr.UnknownValueError:
                # Try with OpenAI Whisper if available
                return await self._transcribe_with_openai(file_path)
        except Exception as e:
            logger.error(f"Audio transcription error for {file_path}: {e}")
            return ""
    
    async def _transcribe_with_openai(self, file_path: str) -> str:
        """Transcribe audio using OpenAI Whisper"""
        try:
            # This would require OpenAI audio transcription API
            # For now, return empty string
            logger.warning("OpenAI audio transcription not implemented yet")
            return ""
        except Exception as e:
            logger.error(f"OpenAI transcription error: {e}")
            return ""
    
    async def _create_chunks_with_embeddings(self, document: KnowledgeDocument) -> list[KnowledgeChunk]:
        """Create chunks and generate embeddings"""
        chunks = []
        text_chunks = self.qdrant_service._chunk_text(document.content)
        
        for i, chunk_text in enumerate(text_chunks):
            # Generate embedding
            embedding = await self.openai_service.generate_embedding(chunk_text)
            
            # Create chunk metadata
            chunk_metadata = {
                "document_id": document.document_id,
                "filename": document.filename,
                "file_type": document.file_type,
                "chunk_index": i,
                "total_chunks": len(text_chunks),
                **document.metadata
            }
            
            chunk = KnowledgeChunk(
                chunk_id=f"{document.document_id}_chunk_{i}",
                document_id=document.document_id,
                content=chunk_text,
                embedding=embedding,
                metadata=chunk_metadata,
                created_at=datetime.now()
            )
            chunks.append(chunk)
        
        return chunks
    
    async def _update_upload_status(self, upload_id: str, status: str, progress: float, error_message: str = None):
        """Update upload status (this would update database in production)"""
        # For now, this is a placeholder
        # In production, you'd update a database or Redis
        logger.info(f"Upload {upload_id}: {status} - {progress:.1%}")
    
    async def delete_file_data(self, upload_id: str):
        """Delete file and its processed data"""
        try:
            # Delete from Qdrant
            await self.qdrant_service.delete_document(upload_id)
            
            # Delete physical file
            # Note: This would require tracking the actual file path
            logger.info(f"Deleted data for upload: {upload_id}")
            
        except Exception as e:
            logger.error(f"Error deleting file data {upload_id}: {e}")
            raise
